import os
import docx
from docx.shared import Pt

class DocxGeneratorTool:
    @staticmethod
    def create_document(filename: str, title: str, content: str) -> str:
        os.makedirs("./data/reports", exist_ok=True)
        file_path = f"./data/reports/{filename}"
        
        doc = docx.Document()
        doc.add_heading(title, level=0)
        
        for para in content.split("\n"):
            if para.strip():
                p = doc.add_paragraph()
                run = p.add_run(para.strip())
                run.font.size = Pt(11)
                
        doc.save(file_path)
        return os.path.abspath(file_path)
